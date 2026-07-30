from docx import Document
from docx.shared import Pt


def export_docx(paragraphs, output_path, title="강의 전사본"):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"

    doc.add_heading(title, level=1)
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(output_path)